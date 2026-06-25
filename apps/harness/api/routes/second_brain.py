from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from core.auth import require_auth
from db.session import get_db
from db.models import KnowledgeItem, DocumentChunk
from memory import second_brain

router = APIRouter()


# ─── Response models ──────────────────────────────────────────────────────────

class KnowledgeItemOut(BaseModel):
    id: str
    type: str
    url: Optional[str]
    source_title: Optional[str]
    source_author: Optional[str]
    summary: Optional[str]
    personal_note: Optional[str]
    tags: list
    domain: Optional[str]
    access_count: int
    starred: bool = False
    properties: dict = {}
    saved_at: datetime

    class Config:
        from_attributes = True


class KnowledgeItemDetailOut(KnowledgeItemOut):
    """Full detail — includes content and chunk count. Used for the detail panel."""
    clean_content: Optional[str] = None
    chunk_count: int = 0


class IngestUrlRequest(BaseModel):
    url: str
    personal_note: str = ""
    tags: List[str] = []
    domain: Optional[str] = None


class IngestTextRequest(BaseModel):
    content: str
    title: str = ""
    personal_note: str = ""
    tags: List[str] = []
    domain: Optional[str] = None


class UpdateItemRequest(BaseModel):
    source_title: Optional[str] = None
    personal_note: Optional[str] = None
    tags: Optional[List[str]] = None
    domain: Optional[str] = None
    starred: Optional[bool] = None
    properties: Optional[dict] = None
    clean_content: Optional[str] = None   # document content edit — triggers re-embed + re-chunk


class IngestDocumentRequest(BaseModel):
    content: str
    title: str = ""
    personal_note: str = ""
    tags: List[str] = []
    domain: Optional[str] = None


class EnhanceRequest(BaseModel):
    selected_text: str
    action: str  # "improve"|"shorten"|"expand"|"rephrase"|"continue"|"custom"
    custom_prompt: Optional[str] = None
    document_context: Optional[str] = None


class GenerateRequest(BaseModel):
    prompt: str
    document_context: Optional[str] = None   # existing doc text (style/topic context)
    cursor_context: Optional[str] = None      # text just before the cursor


class SearchRequest(BaseModel):
    query: str
    limit: int = 5


class SearchResultOut(BaseModel):
    item_id: str
    item_title: Optional[str]
    item_type: str
    chunk_content: Optional[str]
    url: Optional[str]
    chunk_index: Optional[int] = None


# ─── Routes ───────────────────────────────────────────────────────────────────

@router.get("/items", response_model=List[KnowledgeItemOut])
async def list_items(
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    return await second_brain.list_items(db, user_id)


@router.get("/items/{item_id}", response_model=KnowledgeItemDetailOut)
async def get_item(
    item_id: str,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    item = await second_brain.get_item(db, item_id, user_id)
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")

    # Count chunks
    count_result = await db.execute(
        select(func.count(DocumentChunk.id)).where(
            DocumentChunk.knowledge_item_id == item_id
        )
    )
    chunk_count = count_result.scalar() or 0

    return KnowledgeItemDetailOut(
        id=item.id,
        type=item.type,
        url=item.url,
        source_title=item.source_title,
        source_author=item.source_author,
        summary=item.summary,
        personal_note=item.personal_note,
        tags=item.tags or [],
        domain=item.domain,
        access_count=item.access_count,
        starred=item.starred,
        properties=item.properties or {},
        saved_at=item.saved_at,
        clean_content=item.clean_content,
        chunk_count=chunk_count,
    )


@router.patch("/items/{item_id}", response_model=KnowledgeItemDetailOut)
async def update_item(
    item_id: str,
    body: UpdateItemRequest,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    item = await second_brain.get_item(db, item_id, user_id)
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")

    if body.source_title is not None:
        item.source_title = body.source_title
    if body.personal_note is not None:
        item.personal_note = body.personal_note
    if body.tags is not None:
        item.tags = body.tags
    if body.domain is not None:
        item.domain = body.domain
    if body.starred is not None:
        item.starred = body.starred
    if body.properties is not None:
        item.properties = {**(item.properties or {}), **body.properties}
    if body.clean_content is not None:
        import re as _re
        from sqlalchemy import delete as _delete
        from memory.embeddings import embed_one, embed
        from memory.chunker import chunk_text
        item.clean_content = body.clean_content
        item.raw_content = body.clean_content
        # Strip [[id|type|label]] mention markers before embedding/chunking
        # so the stored round-trip format doesn't pollute semantic search
        clean_for_embed = _re.sub(r'\[\[[^\]|]+\|[^\]|]+\|([^\]]+)\]\]', r'\1', body.clean_content)
        # Re-compute item-level embedding
        summary_text = " ".join(clean_for_embed.split()[:800])
        item.embedding = embed_one(summary_text) if summary_text else None
        item.summary = summary_text[:500] if summary_text else None
        # Delete old chunks and re-chunk
        await db.execute(
            _delete(DocumentChunk).where(DocumentChunk.knowledge_item_id == item_id)
        )
        chunks = chunk_text(clean_for_embed)
        if chunks:
            chunk_embeddings = embed(chunks)
            for i, (chunk_str, emb) in enumerate(zip(chunks, chunk_embeddings)):
                db.add(DocumentChunk(
                    knowledge_item_id=item.id,
                    chunk_index=i,
                    content=chunk_str,
                    embedding=emb,
                    token_count=len(chunk_str.split()),
                ))

    await db.commit()
    await db.refresh(item)

    count_result = await db.execute(
        select(func.count(DocumentChunk.id)).where(
            DocumentChunk.knowledge_item_id == item_id
        )
    )
    chunk_count = count_result.scalar() or 0

    return KnowledgeItemDetailOut(
        id=item.id,
        type=item.type,
        url=item.url,
        source_title=item.source_title,
        source_author=item.source_author,
        summary=item.summary,
        personal_note=item.personal_note,
        tags=item.tags or [],
        domain=item.domain,
        access_count=item.access_count,
        starred=item.starred,
        properties=item.properties or {},
        saved_at=item.saved_at,
        clean_content=item.clean_content,
        chunk_count=chunk_count,
    )


@router.post("/ingest/url", response_model=KnowledgeItemOut, status_code=201)
async def ingest_url(
    body: IngestUrlRequest,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    try:
        item = await second_brain.ingest_url(
            db, user_id, body.url,
            personal_note=body.personal_note,
            tags=body.tags,
            domain=body.domain,
        )
        return item
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Ingest failed: {e}")


@router.post("/ingest/text", response_model=KnowledgeItemOut, status_code=201)
async def ingest_text(
    body: IngestTextRequest,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    item = await second_brain.ingest_text(
        db, user_id, body.content,
        title=body.title,
        personal_note=body.personal_note,
        tags=body.tags,
        domain=body.domain,
    )
    return item


@router.post("/ingest/document", response_model=KnowledgeItemOut, status_code=201)
async def ingest_document(
    body: IngestDocumentRequest,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    item = await second_brain.ingest_document(
        db, user_id, body.content,
        title=body.title,
        personal_note=body.personal_note,
        tags=body.tags,
        domain=body.domain,
    )
    return item


@router.post("/ai/enhance")
async def ai_enhance(
    body: EnhanceRequest,
    user_id: str = Depends(require_auth),
):
    from fastapi.responses import StreamingResponse
    import anthropic as _anthropic
    from core.config import settings
    from core.streaming import sse_event, sse_done

    action_prompts = {
        "improve": "Improve the clarity, flow, and impact of this text. Keep approximately the same length and preserve the core meaning.",
        "shorten": "Shorten this text significantly while preserving the key points.",
        "expand": "Expand this text with more detail, context, or examples. Keep the same style.",
        "rephrase": "Rephrase this text in a different way while keeping the exact same meaning.",
        "continue": "Continue writing from where this text ends. Match the style, tone, and context.",
        "custom": body.custom_prompt or "Rewrite this text to be better.",
    }

    system_prompt = (
        "You are a precise writing assistant. "
        "Return ONLY the replacement text — no preamble, no explanation, no surrounding quotes. "
        "Preserve markdown formatting (bold, headers, lists, etc.) where present."
    )

    context_block = (
        f"\n\nDocument context (for reference only — do not reproduce it):\n{body.document_context[:1500]}"
        if body.document_context else ""
    )
    user_message = (
        f"{action_prompts[body.action]}\n\n"
        f"Text:\n{body.selected_text}"
        f"{context_block}"
    )

    async def generate():
        try:
            from core.model_client import get_model_client as _gmc
            _p1 = settings.tier1_provider
            client = _gmc().zai if _p1 == "zai" else _gmc().anthropic
            _m1 = settings.tier1_model_override or ("glm-4.5-air" if _p1 == "zai" else settings.tier1_model)
            async with client.messages.stream(
                model=_m1,
                max_tokens=1024,
                system=system_prompt,
                messages=[{"role": "user", "content": user_message}],
            ) as stream:
                async for text in stream.text_stream:
                    yield sse_event({"type": "chunk", "text": text})
        except Exception as exc:
            yield sse_event({"type": "error", "message": str(exc)})
        finally:
            yield sse_done()

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/ai/generate")
async def ai_generate(
    body: GenerateRequest,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    """
    Freestyle generation with full TARS context.
    Uses Sonnet + context assembler (memories, calendar, tasks, meetings, second brain).
    Streams SSE chunks.
    """
    from fastapi.responses import StreamingResponse
    import anthropic as _anthropic
    from core.config import settings
    from core.streaming import sse_event, sse_done
    from core import context_assembler
    from core.model_client import ModelTier

    # Build full TARS system prompt (same quality as the chat window)
    system_prompt = await context_assembler.assemble(
        user_id=user_id,
        query=body.prompt,
        db=db,
        tier=ModelTier.TIER3,
    )
    system_prompt += (
        "\n\n[DOCUMENT WRITING MODE]\n"
        "You are writing content directly into a document. "
        "Return ONLY the generated text — no preamble, no meta-commentary, no surrounding quotes. "
        "Use markdown formatting naturally (headers, bold, lists, etc.) where it fits. "
        "Match the style and tone of any existing document content provided."
    )

    # Build the user message
    parts: list[str] = []
    if body.document_context:
        parts.append(f"Existing document content (for style/context — do not repeat it):\n{body.document_context[:2000]}")
    if body.cursor_context:
        parts.append(f"Text immediately before the cursor:\n{body.cursor_context[-600:]}")
    parts.append(f"Write: {body.prompt}")
    user_message = "\n\n".join(parts)

    # Use writing category routing if configured; fall back to tier3.
    _writing_override = settings.category_routing().get("writing")
    if _writing_override:
        _gen_provider = _writing_override["provider"]
        _gen_model = _writing_override["model"]
    else:
        _gen_provider = settings.tier3_provider
        _gen_model = settings.tier3_model_override or ("glm-4.7" if _gen_provider == "zai" else "claude-sonnet-4-6")

    async def generate():
        try:
            from core.model_client import get_model_client as _gmc2
            client = _gmc2().zai if _gen_provider == "zai" else _gmc2().anthropic
            async with client.messages.stream(
                model=_gen_model,
                max_tokens=2048,
                system=system_prompt,
                messages=[{"role": "user", "content": user_message}],
            ) as stream:
                async for text in stream.text_stream:
                    yield sse_event({"type": "chunk", "text": text})
        except Exception as exc:
            yield sse_event({"type": "error", "message": str(exc)})
        finally:
            yield sse_done()

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.post("/search", response_model=List[SearchResultOut])
async def search(
    body: SearchRequest,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    results = await second_brain.search(db, user_id, body.query, limit=body.limit)
    return [
        SearchResultOut(
            item_id=r["item"].id,
            item_title=r["item"].source_title,
            item_type=r["item"].type,
            chunk_content=r["chunk"].content if r["chunk"] else r["item"].summary,
            chunk_index=r["chunk"].chunk_index if r["chunk"] else None,
            url=r["item"].url,
        )
        for r in results
    ]


@router.delete("/items/{item_id}", status_code=204)
async def delete_item(
    item_id: str,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    removed = await second_brain.remove_item(db, item_id, user_id)
    if not removed:
        raise HTTPException(status_code=404, detail="Item not found")


@router.get("/items/{item_id}/export")
async def export_item(
    item_id: str,
    format: str,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    """
    Export a Second Brain item as DOCX, PDF, or Google Doc.
    format: "docx" | "pdf" | "gdoc"
    """
    from fastapi.responses import Response, JSONResponse
    import io
    import asyncio
    import re

    item = await second_brain.get_item(db, item_id, user_id)
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")

    content = item.clean_content or item.summary or ""
    title = item.source_title or "Untitled"

    # Strip [[id|type|label]] mention markers → plain labels
    content = re.sub(r'\[\[[^\]|]+\|[^\]|]+\|([^\]]+)\]\]', r'@\1', content)

    # Shared DOCX builder — used by the docx download and by the Google Doc export
    # (uploaded to Drive with conversion so the Doc is rich text, not literal markdown).
    def build_docx() -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Render inline markdown (**bold**, *italic*, `code`) into runs on a paragraph.
        def add_inline(para, text):
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    para.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    para.add_run(part[1:-1]).italic = True
                elif part.startswith("`") and part.endswith("`"):
                    run = para.add_run(part[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                else:
                    para.add_run(part)

        # Document title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Personal note as italic block if present
        if item.personal_note:
            note_para = doc.add_paragraph()
            note_run = note_para.add_run(f"Note: {item.personal_note}")
            note_run.italic = True
            note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()

        # Parse markdown line-by-line into docx
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            # Headings
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            # Unordered list items — inline-parsed
            elif re.match(r'^[-*+] ', line):
                add_inline(doc.add_paragraph(style="List Bullet"), line[2:])
            # Ordered list items — inline-parsed
            elif re.match(r'^\d+\. ', line):
                add_inline(doc.add_paragraph(style="List Number"), re.sub(r'^\d+\. ', '', line))
            # Horizontal rule
            elif line.strip() in ("---", "***", "___"):
                doc.add_paragraph("─" * 40)
            # Blank line
            elif line.strip() == "":
                pass
            else:
                # Normal paragraph — inline **bold**, *italic*, `code`
                add_inline(doc.add_paragraph(), line)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    if format == "docx":
        loop = asyncio.get_running_loop()
        docx_bytes = await loop.run_in_executor(None, build_docx)
        safe_name = re.sub(r'[^\w\s-]', '', title)[:60].strip().replace(' ', '_') or 'export'
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )

    elif format == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_LEFT
        import re as _re

        def build_pdf() -> bytes:
            buf = io.BytesIO()
            doc_pdf = SimpleDocTemplate(
                buf,
                pagesize=A4,
                leftMargin=25 * mm,
                rightMargin=25 * mm,
                topMargin=25 * mm,
                bottomMargin=25 * mm,
            )
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                "TARSTitle",
                parent=styles["Heading1"],
                fontSize=20,
                leading=26,
                spaceAfter=6,
                textColor=colors.HexColor("#1a1a1a"),
            )
            h1_style = ParagraphStyle("TARSH1", parent=styles["Heading1"], fontSize=16, leading=20, spaceAfter=4, spaceBefore=10)
            h2_style = ParagraphStyle("TARSH2", parent=styles["Heading2"], fontSize=13, leading=17, spaceAfter=3, spaceBefore=8)
            h3_style = ParagraphStyle("TARSH3", parent=styles["Heading3"], fontSize=11, leading=15, spaceAfter=2, spaceBefore=6)
            body_style = ParagraphStyle("TARSBody", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=4)
            note_style = ParagraphStyle("TARSNote", parent=styles["Normal"], fontSize=9, leading=13, textColor=colors.HexColor("#666666"), fontName="Helvetica-Oblique")
            bullet_style = ParagraphStyle("TARSBullet", parent=styles["Normal"], fontSize=10, leading=14, leftIndent=12, spaceAfter=2, bulletIndent=0)
            code_style = ParagraphStyle("TARSCode", parent=styles["Code"], fontSize=9, leading=12, backColor=colors.HexColor("#f4f4f4"))

            def inline_md(text):
                """Convert inline **bold** and *italic* to reportlab markup."""
                text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                text = _re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                text = _re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
                text = _re.sub(r'`(.+?)`', r'<font name="Courier">\1</font>', text)
                return text

            story = []
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 4 * mm))

            if item.personal_note:
                story.append(Paragraph(f"Note: {item.personal_note}", note_style))
                story.append(Spacer(1, 3 * mm))

            lines = content.split("\n")
            for line in lines:
                if line.startswith("### "):
                    story.append(Paragraph(inline_md(line[4:]), h3_style))
                elif line.startswith("## "):
                    story.append(Paragraph(inline_md(line[3:]), h2_style))
                elif line.startswith("# "):
                    story.append(Paragraph(inline_md(line[2:]), h1_style))
                elif _re.match(r'^[-*+] ', line):
                    story.append(Paragraph(f"• {inline_md(line[2:])}", bullet_style))
                elif _re.match(r'^\d+\. ', line):
                    num = _re.match(r'^(\d+)\. ', line).group(1)
                    line_text = _re.sub(r'^\d+\. ', '', line)
                    story.append(Paragraph(f"{num}. {inline_md(line_text)}", bullet_style))
                elif line.strip() in ("---", "***", "___"):
                    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cccccc"), spaceAfter=4))
                elif line.strip() == "":
                    story.append(Spacer(1, 3 * mm))
                else:
                    story.append(Paragraph(inline_md(line), body_style))

            doc_pdf.build(story)
            buf.seek(0)
            return buf.read()

        loop = asyncio.get_running_loop()
        pdf_bytes = await loop.run_in_executor(None, build_pdf)
        safe_name = re.sub(r'[^\w\s-]', '', title)[:60].strip().replace(' ', '_') or 'export'
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
        )

    elif format == "gdoc":
        from memory.second_brain import load_workspace_client

        client = await load_workspace_client(db)
        if client is None:
            raise HTTPException(
                status_code=503,
                detail="Google Workspace is not connected. Connect it in Connectors first.",
            )

        loop = asyncio.get_running_loop()
        # Build the same DOCX as the download path, then upload to Drive with
        # conversion so the Google Doc is rich text (headings, bold, lists) rather
        # than literal markdown characters.
        docx_bytes = await loop.run_in_executor(None, build_docx)
        result = await loop.run_in_executor(
            None, client.create_doc_from_docx, title, docx_bytes
        )
        return JSONResponse({"url": result["url"], "title": result["title"]})

    else:
        raise HTTPException(status_code=400, detail=f"Unknown format: {format}. Use docx, pdf, or gdoc.")


@router.post("/items/{item_id}/properties/auto", response_model=KnowledgeItemDetailOut)
async def auto_populate_properties(
    item_id: str,
    user_id: str = Depends(require_auth),
    db: AsyncSession = Depends(get_db),
):
    """Use Haiku to infer status/type/priority from item content and merge into properties."""
    item = await second_brain.get_item(db, item_id, user_id)
    if not item:
        raise HTTPException(status_code=404, detail="Item not found")

    from core.config import settings as _cfg
    from core.model_client import get_model_client as _gmc

    content_snippet = (item.summary or item.clean_content or "")[:1200]
    title = item.source_title or ""

    prompt = (
        f'Analyze this knowledge item and return ONLY a valid JSON object with these exact keys:\n'
        f'{{"status": "<raw|developing|actionable|archived>", '
        f'"type": "<idea|project|reference|research>", '
        f'"priority": "<high|medium|low>"}}\n\n'
        f'Title: {title}\n'
        f'Content: {content_snippet}\n\n'
        f'Choose the values that best fit. Return ONLY the JSON object, no explanation.'
    )

    try:
        _p = _cfg.tier1_provider
        client = _gmc().zai if _p == "zai" else _gmc().anthropic
        _model = _cfg.tier1_model_override or ("glm-4.5-air" if _p == "zai" else _cfg.tier1_model)
        resp = await client.messages.create(
            model=_model,
            max_tokens=64,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = resp.content[0].text.strip()
        import json as _json
        # Strip markdown fences if present
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.strip()
        inferred = _json.loads(raw)
    except Exception:
        inferred = {}

    if inferred:
        item.properties = {**(item.properties or {}), **inferred}
        await db.commit()
        await db.refresh(item)

    count_result = await db.execute(
        select(func.count(DocumentChunk.id)).where(
            DocumentChunk.knowledge_item_id == item_id
        )
    )
    chunk_count = count_result.scalar() or 0

    return KnowledgeItemDetailOut(
        id=item.id,
        type=item.type,
        url=item.url,
        source_title=item.source_title,
        source_author=item.source_author,
        summary=item.summary,
        personal_note=item.personal_note,
        tags=item.tags or [],
        domain=item.domain,
        access_count=item.access_count,
        starred=item.starred,
        properties=item.properties or {},
        saved_at=item.saved_at,
        clean_content=item.clean_content,
        chunk_count=chunk_count,
    )
